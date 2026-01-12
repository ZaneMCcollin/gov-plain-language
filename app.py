# ============================================================
# GovCan Plain Language Converter — ONE-FILE ENTERPRISE BUILD
# (Cloud Run + SQLite + Reviewer workflow + Versioning + DOCX/PDF + OCR + Limits + Lang detect + OCR confidence)
#
# This single app.py:
# - Writes requirements.txt / Dockerfile / .dockerignore / .streamlit/config.toml if missing
# - Uses SQLite persistence (versions + approvals + comments persist)
# - Saves versions + approvals + comments (persist)
# - Rollback to any version
# - DOCX + PDF export (no LibreOffice)
# - OCR fallback for scanned PDFs + clear OCR indicators
# - OCR confidence highlighting in UI + confidence stats saved to meta
# - Auto language detection (input + OCR) to guide OCR + Gemini prompts
# - Word/char limits BEFORE LLM calls to reduce 429/cost
# - Streamlit Cloud ready (no fixed port in config.toml; Docker uses $PORT)
# - Watermark: DRAFT vs APPROVED on PDFs + status banner on DOCX
# - Compliance Report export (PDF summary)
# - ✅ Auth: Streamlit built-in login (Google/OIDC) + domain/email allowlist
# - ✅ Usage analytics: per-user events for billing insight + admin CSV export
# ============================================================

import io
import re
import json
import time
import csv
import hashlib
import sqlite3
import sys
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st

# Auth/roles helpers (separated module)
from auth import require_login as auth_require_login, get_effective_role as auth_get_effective_role
import os

def ensure_secrets_toml_from_env() -> None:
    """On Cloud Run, write Streamlit secrets.toml from environment variables.

    Why: st.secrets is file-based. Cloud Run typically provides env vars.
    This keeps one codebase that works on both Streamlit Community Cloud and Cloud Run.

    Notes:
    - We ALWAYS (re)write secrets.toml on Cloud Run so stale/empty files don't break auth.
    - We write to both /app/.streamlit and /root/.streamlit (Streamlit checks both).
    """
    if not os.environ.get("K_SERVICE"):
        return  # not Cloud Run

    gemini = (os.environ.get("GEMINI_API_KEY") or "").strip()
    if not gemini:
        # Don't create an empty secrets file; the app will show a clear error later.
        return

    lines: list[str] = []

    def add(s: str) -> None:
        lines.append(str(s))

    # Top-level keys
    add(f'GEMINI_API_KEY = "{gemini}"')

    for k in ("PROD", "DEBUG", "ALLOWED_EMAILS", "ALLOWED_DOMAINS", "ENABLE_ROLE_SWITCH", "ENABLE_WORKSPACE_SWITCH"):
        v = (os.environ.get(k) or "").strip()
        if v:
            add(f'{k} = "{v}"')

    # Auth (Streamlit built-in auth)
    auth_redirect = (os.environ.get("AUTH_REDIRECT_URI") or "").strip()
    auth_cookie = (os.environ.get("AUTH_COOKIE_SECRET") or "").strip()
    auth_meta = (os.environ.get("AUTH_SERVER_METADATA_URL") or "https://accounts.google.com/.well-known/openid-configuration").strip()
    google_id = (os.environ.get("AUTH_GOOGLE_CLIENT_ID") or "").strip()
    google_secret = (os.environ.get("AUTH_GOOGLE_CLIENT_SECRET") or "").strip()

    # Only write auth section if we have the minimum needed.
    if auth_redirect and auth_cookie and google_id and google_secret:
        add("")
        add("[auth]")
        add(f'redirect_uri = "{auth_redirect}"')
        add(f'cookie_secret = "{auth_cookie}"')
        add(f'server_metadata_url = "{auth_meta}"')
        add("")
        add("[auth.google]")
        add(f'client_id = "{google_id}"')
        add(f'client_secret = "{google_secret}"')
        add("")

    payload = "\n".join(lines) + "\n"

    for path in ("/app/.streamlit/secrets.toml", "/root/.streamlit/secrets.toml"):
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(payload)


# MUST be called immediately after definition
ensure_secrets_toml_from_env()



from google import genai
from google.genai.errors import ClientError

from pypdf import PdfReader
from docx import Document

import pytesseract
import pypdfium2 as pdfium
from PIL import Image  # noqa: F401

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

# ============================================================
# Auto language detection (patched to show real error)
# ============================================================
LANGDETECT_OK = False
LANGDETECT_ERR = ""

try:
    from langdetect import detect, DetectorFactory
    DetectorFactory.seed = 0
    LANGDETECT_OK = True
except Exception as e:
    LANGDETECT_OK = False
    LANGDETECT_ERR = repr(e)

# ============================================================
# Project files (auto-create if missing)
# ============================================================
def _write_if_missing(path: str, content: str) -> None:
    if os.path.exists(path):
        return
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)

def _bootstrap_project_files() -> None:
    # ✅ Pin Streamlit new enough to support st.login / st.user
    # ✅ Authlib required for Google auth on Streamlit Community Cloud
    _write_if_missing(
        "requirements.txt",
        "\n".join([
            "streamlit>=1.42.0",
            "Authlib>=1.3.2",
            "google-genai",
            "pypdf",
            "python-docx",
            "pytesseract",
            "pypdfium2",
            "Pillow",
            "reportlab",
            "langdetect",
            ""
        ])
    )

    _write_if_missing(
        "Dockerfile",
        "\n".join([
            "FROM python:3.11-slim",
            "",
            "ENV PYTHONDONTWRITEBYTECODE=1",
            "ENV PYTHONUNBUFFERED=1",
            "",
            "# System deps for OCR (English + French + OSD)",
            "RUN apt-get update && apt-get install -y --no-install-recommends \\",
            "    tesseract-ocr \\",
            "    tesseract-ocr-eng \\",
            "    tesseract-ocr-fra \\",
            "    tesseract-ocr-osd \\",
            "    && rm -rf /var/lib/apt/lists/*",
            "",
            "WORKDIR /app",
            "",
            "COPY requirements.txt /app/requirements.txt",
            "RUN pip install --no-cache-dir -r requirements.txt",
            "",
            "COPY app.py /app/app.py",
            "",
            "ENV PORT=8080",
            "EXPOSE 8080",
            "",
            'CMD ["bash","-lc","streamlit run app.py --server.port=${PORT:-8080} --server.address=0.0.0.0 --server.headless=true --server.enableCORS=false --server.enableXsrfProtection=false"]',
            ""
        ])
    )

    _write_if_missing(
        ".dockerignore",
        "\n".join([
            ".venv",
            "__pycache__",
            "*.pyc",
            ".DS_Store",
            ".git",
            ".gitignore",
            ".env",
            ".streamlit",
            "data",
            ""
        ])
    )

    # ✅ Streamlit Cloud friendly: do NOT pin port/address here
    os.makedirs(".streamlit", exist_ok=True)
    _write_if_missing(
        ".streamlit/config.toml",
        "\n".join([
            "[server]",
            "headless = true",
            "enableCORS = false",
            "enableXsrfProtection = false",
            "",
            "[browser]",
            "gatherUsageStats = false",
            ""
        ])
    )

    # ✅ Minimal unit tests (stdlib unittest; no extra deps)
    _write_if_missing(
        os.path.join("tests", "test_grading.py"),
        "\n".join([
            "import unittest",
            "",
            "# Import the grading module (streamlit-free)",
            "import grading as g",
            "",
            "",
            "class TestReadability(unittest.TestCase):",
            "    def test_flesch_kincaid_simple_is_reasonable(self):",
            "        text = 'This is a simple sentence. It is easy to read.'",
            "        grade = g.flesch_kincaid(text)",
            "        # Not asserting an exact grade (heuristic), only that it's not extreme.",
            "        self.assertGreaterEqual(grade, 0.0)",
            "        self.assertLessEqual(grade, 12.0)",
            "",
            "    def test_flesch_kincaid_complex_harder_than_simple(self):",
            "        simple = 'We will send the form today.'",
            "        complex_ = ('Pursuant to the aforementioned regulatory framework, the department will '",
            "                    'expeditiously disseminate the requisite documentation at its earliest convenience.')",
            "        self.assertGreater(g.flesch_kincaid(complex_), g.flesch_kincaid(simple))",
            "",
            "    def test_sentence_split(self):",
            "        text = 'One. Two! Three? Four'",
            "        parts = g.split_sentences(text)",
            "        self.assertEqual(len(parts), 4)",
            "",
            "",
            "if __name__ == '__main__':",
            "    unittest.main()",
            ""
        ])
    )

_bootstrap_project_files()

# ============================================================
# Runtime configuration
# ============================================================
MAX_DOC_CHARS = int(os.environ.get("MAX_DOC_CHARS", "80000"))
MAX_CHUNK_WORDS = int(os.environ.get("MAX_CHUNK_WORDS", "900"))
LLM_RETRIES = int(os.environ.get("LLM_RETRIES", "5"))
LLM_MODEL = os.environ.get("LLM_MODEL", "gemini-2.0-flash")

LOCAL_DATA_DIR = os.environ.get("LOCAL_DATA_DIR", "data").strip()
SQLITE_PATH = os.environ.get("SQLITE_PATH", os.path.join(LOCAL_DATA_DIR, "app.db")).strip()

OCR_LOW_CONF = int(os.environ.get("OCR_LOW_CONF", "60"))
OCR_MED_CONF = int(os.environ.get("OCR_MED_CONF", "80"))

BILLING_RATE_PER_1K = float(os.environ.get("BILLING_RATE_PER_1K", "0") or "0")
BILLING_MODEL = (os.environ.get("BILLING_MODEL", "per_doc") or "per_doc").strip().lower()
BILLING_PRICE_PER_DOC = float(os.environ.get("BILLING_PRICE_PER_DOC", os.environ.get("BILLING_PRICE_PER_DOC_CAD", "25")) or "25")

# ============================================================
# ENV FLAGS (define PROD before using it)
# ============================================================

def safe_secret(key: str, default=None):
    """Return a config value from env first, then Streamlit secrets."""
    v = os.environ.get(key)
    if v is not None and str(v).strip() != "":
        return v
    try:
        return st.secrets.get(key, default)
    except Exception:
        return default

# PROD: true on Streamlit Cloud unless explicitly overridden
PROD = str(os.environ.get("PROD", "") or safe_secret("PROD", "") or "").lower() in ("1", "true", "yes")

# DEBUG: only allowed when NOT PROD, and only if explicitly enabled
DEBUG = (not PROD) and (
    str(os.environ.get("DEBUG", "") or safe_secret("DEBUG", "false") or "").lower()
    in ("1", "true", "yes")
)

# ============================================================
# Page config
# ============================================================
st.set_page_config(page_title="GovCan Plain Language Converter", layout="wide")
st.title("🇨🇦 GovCan Plain Language Converter")

# ============================================================
# Hosted-mode warning (Streamlit Cloud demo: disk persistence not guaranteed)
# ============================================================
def _is_streamlit_cloud() -> bool:
    # Avoid false-positives on Cloud Run (we also run headless there).
    if os.environ.get("K_SERVICE"):
        return False
    return bool(
        os.environ.get("STREAMLIT_SHARING")
        or os.environ.get("STREAMLIT_CLOUD")
        or os.environ.get("STREAMLIT_COMMUNITY_CLOUD")
    )


# Resolve once (used by auth)
IS_STREAMLIT_CLOUD = _is_streamlit_cloud()

if IS_STREAMLIT_CLOUD:
    st.caption("Running in hosted mode. Note: local file storage (including SQLite) may not persist across reboots on Streamlit Community Cloud.")

# ============================================================
# ============================================================
# ✅ Authentication + Allow-lists + Locked roles (from Secrets)
# ============================================================

from collections.abc import Mapping

def _get_allowlists() -> Tuple[List[str], List[str]]:
    """Read allow-lists from Secrets (comma-separated strings)."""
    allowed_domains = safe_secret("ALLOWED_DOMAINS", "")
    allowed_emails = safe_secret("ALLOWED_EMAILS", "")
    doms = [d.strip().lower() for d in str(allowed_domains).split(",") if d.strip()]
    ems = [e.strip().lower() for e in str(allowed_emails).split(",") if e.strip()]
    return doms, ems

def _normalize_email_list(x) -> List[str]:
    if not x:
        return []
    if isinstance(x, str):
        return [e.strip().lower() for e in x.split(",") if e.strip()]
    if isinstance(x, (list, tuple)):
        return [str(e).strip().lower() for e in x if str(e).strip()]
    return []

def _roles_config() -> Dict[str, List[str]]:
    """
    Reads roles from config.

    Supported sources:
    - Streamlit secrets: [roles] table
    - Cloud Run / env: ROLES (JSON) or roles (JSON) or ROLES_ADMIN/ROLES_EDITOR/...

    Examples (JSON env):
      roles={"admin":["a@b.com"],"editor":["e@b.com"],"reviewer":[],"viewer":[]}

    Secrets format:
      [roles]
      admin = "a@b.com"
      reviewer = "c@d.com"
      editor = "e@d.com,f@g.com"
      viewer = "h@i.com"
    """
    r = safe_secret("roles", None)

    # --- If roles came from env, it's often a JSON string. Try to parse it.
    if isinstance(r, str) and r.strip():
        try:
            parsed = json.loads(r)
            r = parsed
        except Exception:
            # If not JSON, treat it as empty and fall back to ROLES_* env vars below.
            r = None

    # --- If roles is not a mapping (e.g., missing), allow ROLES_* env vars as a fallback.
    if not isinstance(r, Mapping):
        # Support separate env vars (handy on Cloud Run UI):
        # ROLES_ADMIN="a@b.com,b@c.com", ROLES_EDITOR="...", etc.
        admin = safe_secret("ROLES_ADMIN", "") or ""
        reviewer = safe_secret("ROLES_REVIEWER", "") or ""
        editor = safe_secret("ROLES_EDITOR", "") or ""
        viewer = safe_secret("ROLES_VIEWER", "") or ""
        return {
            "admin": _normalize_email_list(admin),
            "reviewer": _normalize_email_list(reviewer),
            "editor": _normalize_email_list(editor),
            "viewer": _normalize_email_list(viewer),
        }

    return {
        "admin": _normalize_email_list(r.get("admin")),
        "reviewer": _normalize_email_list(r.get("reviewer")),
        "editor": _normalize_email_list(r.get("editor")),
        "viewer": _normalize_email_list(r.get("viewer")),
    }

def _superadmin_emails() -> List[str]:
    """Emails that should ALWAYS be admin, regardless of roles config."""
    # Allow both SINGLE and list versions (env or secrets)
    one = safe_secret("SUPERADMIN_EMAIL", "") or ""
    many = safe_secret("SUPERADMIN_EMAILS", "") or ""
    out: List[str] = []
    out += _normalize_email_list(one)
    out += _normalize_email_list(many)
    # de-dupe
    return sorted({e for e in out if e})

def _ensure_role_tables(conn: sqlite3.Connection) -> None:
    """Ensure role tables exist. Safe to call early (before full _db_init)."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS role_assignments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ts TEXT NOT NULL,
            workspace TEXT NOT NULL,      -- '*' means global
            email TEXT NOT NULL,
            role TEXT NOT NULL,           -- admin/reviewer/editor/viewer
            updated_by TEXT
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_roles_email_ws ON role_assignments(email, workspace)")
    conn.commit()

def _db_role_lookup(email: str, workspace: str) -> Optional[str]:
    """Return most recent role for (email, workspace) or global ('*')."""
    email = (email or "").strip().lower()
    ws = (workspace or "").strip() or "default"
    if not email:
        return None
    try:
        os.makedirs(os.path.dirname(SQLITE_PATH) or ".", exist_ok=True)
        conn = sqlite3.connect(SQLITE_PATH, check_same_thread=False)
        _ensure_role_tables(conn)
        cur = conn.cursor()

        cur.execute(
            "SELECT role FROM role_assignments WHERE email = ? AND workspace = ? ORDER BY id DESC LIMIT 1",
            (email, ws),
        )
        row = cur.fetchone()
        if row and row[0]:
            conn.close()
            return str(row[0])

        cur.execute(
            "SELECT role FROM role_assignments WHERE email = ? AND workspace = '*' ORDER BY id DESC LIMIT 1",
            (email,),
        )
        row = cur.fetchone()
        conn.close()
        return str(row[0]) if row and row[0] else None
    except Exception:
        return None

def _set_role_assignment(workspace: str, email: str, role: str, updated_by: str = "") -> None:
    """Persist a role assignment (safe; never raises)."""
    ws = (workspace or "").strip() or "default"
    em = (email or "").strip().lower()
    rl = (role or "").strip().lower() or "viewer"
    if not em:
        return
    try:
        os.makedirs(os.path.dirname(SQLITE_PATH) or ".", exist_ok=True)
        conn = sqlite3.connect(SQLITE_PATH, check_same_thread=False)
        _ensure_role_tables(conn)
        conn.execute(
            "INSERT INTO role_assignments(ts, workspace, email, role, updated_by) VALUES(?,?,?,?,?)",
            (datetime.now(timezone.utc).isoformat(), ws, em, rl, (updated_by or "").strip().lower()),
        )
        conn.commit()
        conn.close()
    except Exception:
        return

def _list_role_assignments(limit: int = 50) -> List[Dict[str, Any]]:
    """Return latest role assignments for admin UI."""
    try:
        os.makedirs(os.path.dirname(SQLITE_PATH) or ".", exist_ok=True)
        conn = sqlite3.connect(SQLITE_PATH, check_same_thread=False)
        _ensure_role_tables(conn)
        cur = conn.cursor()
        cur.execute(
            "SELECT ts, workspace, email, role, updated_by FROM role_assignments ORDER BY id DESC LIMIT ?",
            (int(limit),),
        )
        rows = cur.fetchall() or []
        conn.close()
        return [
            {"ts": r[0], "workspace": r[1], "email": r[2], "role": r[3], "updated_by": r[4]}
            for r in rows
        ]
    except Exception:
        return []

def role_for_email(email: str) -> str:
    """Back-compat wrapper: returns role string for the given email.
    Real resolution is handled by auth.py (supports SUPERADMIN + config roles + optional DB lookup).
    """
    email = (email or "").strip().lower()
    ws = str(st.session_state.get("workspace", "default") or "default")
    eff = auth_get_effective_role(
        email=email,
        workspace=ws,
        safe_secret=safe_secret,
        prod=PROD,
        db_lookup=_db_role_lookup,
    )
    # Persist helper flags for UI/permissions
    st.session_state["is_global_admin"] = bool(eff.get("is_global_admin"))
    st.session_state["break_glass_admin"] = bool(eff.get("break_glass_admin"))
    st.session_state["role_source"] = str(eff.get("source") or "")
    return str(eff.get("role") or "viewer")

def _user_email() -> str:
    try:
        u = getattr(st, "user", None)
        if not u:
            return ""
        return (getattr(u, "email", "") or "").strip().lower()
    except Exception:
        return ""

def require_login() -> str:
    """Return the logged-in user's email (or "" if not logged in).

    Delegates to auth.py for Streamlit auth + allowlist fallback.
    """
    return auth_require_login(prod=PROD, safe_secret=safe_secret, is_streamlit_cloud=IS_STREAMLIT_CLOUD)

# ============================================================
# Resolve authenticated user
# ============================================================
def log_audit(
    event: str,
    user_email: str = "",
    doc_id: str = "",
    workspace: str = "",
    role: str = "",  # ✅ add this
    meta: Optional[Dict[str, Any]] = None,
) -> None:
    """Write an audit log event to SQLite. Safe: never raises; no-op until DB is ready."""
    try:
        # If DB helpers aren't defined yet at import time, skip safely.
        if "_db" not in globals():
            return

        if not role:
            role = str(st.session_state.get("auth_role", "") or "")

        conn = _db()
        meta_json = json.dumps(meta or {}, ensure_ascii=False)
        ts = datetime.now(timezone.utc).isoformat()

        conn.execute(
            "INSERT INTO audit_logs(ts, user_email, role, event, workspace, doc_id, meta_json) VALUES(?,?,?,?,?,?,?)",
            (ts, (user_email or "").lower(), (role or ""), event, workspace or "", doc_id or "", meta_json),
        )
        conn.commit()
        conn.close()
    except Exception:
        return

AUTH_EMAIL = require_login() or ""
st.session_state["auth_email"] = AUTH_EMAIL

# Compute effective role once per session
if "auth_role" not in st.session_state:
    st.session_state["auth_role"] = role_for_email(AUTH_EMAIL)


# Break-glass audit (SUPERADMIN_EMAIL(S) override)
if st.session_state.get("break_glass_admin") and not st.session_state.get("_break_glass_logged"):
    log_audit(
        event="break_glass_admin",
        user_email=AUTH_EMAIL,
        workspace=str(st.session_state.get("workspace", "default") or "default"),
        doc_id="",
        role="admin",
        meta={"source": st.session_state.get("role_source", "")},
    )
    st.session_state["_break_glass_logged"] = True

# ============================================================
# ðŸ’¼ Client workspaces (multi-tenant namespace)
# ============================================================

def _safe_workspace_key(k: str) -> str:
    k = (k or "").strip().lower()
    k = re.sub(r"[^a-z0-9_-]+", "_", k).strip("_")
    return k or "default"

def _workspaces_config() -> Dict[str, Dict[str, Any]]:
    """Read workspaces from Secrets.

    Example:

    ENABLE_WORKSPACE_SWITCH = "true"

    [workspaces]
    default = { name="Default", domains="", emails="" }
    clienta  = { name="Client A", domains="clienta.com", emails="" }
    clientb  = { name="Client B", domains="", emails="person@clientb.ca" }
    """
    ws = safe_secret("workspaces", {})
    if not isinstance(ws, Mapping):
        return {"default": {"name": "Default", "domains": [], "emails": []}}

    out: Dict[str, Dict[str, Any]] = {}
    for raw_key, cfg in ws.items():
        key = _safe_workspace_key(str(raw_key))
        if not isinstance(cfg, Mapping):
            cfg = {"name": str(raw_key), "domains": "", "emails": ""}
        name = str(cfg.get("name", raw_key))
        domains = _normalize_email_list(cfg.get("domains", ""))  # treat as comma list
        # domains should be pure domains; strip anything after @
        domains = [d.split("@", 1)[-1].lower() for d in domains if d]
        emails = _normalize_email_list(cfg.get("emails", ""))
        out[key] = {"name": name, "domains": domains, "emails": emails}

    if "default" not in out:
        out["default"] = {"name": "Default", "domains": [], "emails": []}
    return out

def workspace_for_email(email: str) -> str:
    email = (email or "").strip().lower()
    dom = email.split("@", 1)[1] if ("@" in email) else ""
    ws = _workspaces_config()

    # Priority: explicit email match, then domain match, else default
    for k, cfg in ws.items():
        if email and email in (cfg.get("emails") or []):
            return k
    for k, cfg in ws.items():
        if dom and dom in (cfg.get("domains") or []):
            return k
    return "default"

def scoped_doc_id(doc_id: str, workspace: str) -> str:
    """Create a storage key that isolates docs per-client."""
    d = (doc_id or "").strip()
    w = _safe_workspace_key(workspace)
    if not d:
        return ""
    return f"{w}::{d}"

# Determine workspace (locked from Secrets) + optional admin override
locked_workspace = workspace_for_email(AUTH_EMAIL) if AUTH_EMAIL else "default"
st.session_state.workspace_locked = locked_workspace

ENABLE_WORKSPACE_SWITCH = str(safe_secret("ENABLE_WORKSPACE_SWITCH", "false")).lower() in ("1", "true", "yes")

active_workspace = locked_workspace

# PROD lock: only global admins can switch workspaces.
if ENABLE_WORKSPACE_SWITCH and st.session_state.get("auth_role") == "admin" and bool(st.session_state.get("is_global_admin")):
    ws_cfg = _workspaces_config()
    ws_keys = sorted(ws_cfg.keys())
    st.sidebar.divider()
    st.sidebar.subheader("Admin: Workspace override (testing)")

    if "workspace_override" not in st.session_state:
        st.session_state.workspace_override = locked_workspace

    st.sidebar.selectbox("Act in workspace", ws_keys, key="workspace_override")

    if st.sidebar.button("Reset workspace to locked", use_container_width=True):
        st.session_state.pop("workspace_override", None)
        st.rerun()

    active_workspace = st.session_state.get("workspace_override", locked_workspace)

st.session_state.workspace = _safe_workspace_key(active_workspace)

# ============================================================
# Auth roles / permissions
# ============================================================

ROLE_PERMS = {
    "admin":    {"convert", "export", "approve", "edit_outputs", "rollback", "analytics"},
    "reviewer": {"export", "approve", "edit_outputs"},
    "editor":   {"convert", "export", "edit_outputs"},
    "viewer":   set(),
}

def can(action: str) -> bool:
    role = st.session_state.get("auth_role", "viewer")
    prod = str(os.getenv("PROD", "false")).lower() == "true"

    # Admin can do everything
    if role == "admin":
        return True

    # Harder PROD lock: only admin can use editor/admin-only capabilities
    if prod and action in {"convert", "edit_outputs", "rollback", "analytics", "role_admin", "workspace_switch"}:
        return False

    return action in ROLE_PERMS.get(role, set())


# ============================================================
# Gemini client

# ============================================================
api_key = safe_secret("GEMINI_API_KEY", "")
if not api_key:
    st.error("❌ GEMINI_API_KEY is missing. Set it in Cloud Run env vars (recommended) or Streamlit Secrets.")
    st.stop()

client = genai.Client(api_key=api_key)

# ============================================================
# Helpers
# ============================================================
def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

def sha12(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()[:12]

def clamp_text(text: str, max_chars: int) -> Tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars], True

def word_count(text: str) -> int:
    return len(re.findall(r"\w+", text))

def truncate_words(text: str, max_words: int) -> Tuple[str, bool]:
    words = text.split()
    if len(words) <= max_words:
        return text, False
    return " ".join(words[:max_words]), True

def safe_filename(name: str) -> str:
    name = re.sub(r"[^a-zA-Z0-9._-]+", "_", name).strip("_")
    return name or "file"

def detect_lang(text: str) -> str:
    if not LANGDETECT_OK:
        return "unknown"
    t = (text or "").strip()
    if len(t) < 40:
        return "unknown"
    try:
        code = detect(t)
        if code.startswith("en"):
            return "en"
        if code.startswith("fr"):
            return "fr"
        return code or "unknown"
    except Exception:
        return "unknown"

def est_tokens_from_text(text: str) -> int:
    if not text:
        return 0
    return int(max(1, len(text) / 4))

def money_estimate(tokens_total: int) -> float:
    if BILLING_RATE_PER_1K <= 0:
        return 0.0
    return round((tokens_total / 1000.0) * BILLING_RATE_PER_1K, 6)

def analytics_summary(days: int = 30) -> Dict[str, Any]:
    since = (datetime.now(timezone.utc).timestamp() - days * 86400)
    since_iso = datetime.fromtimestamp(since, tz=timezone.utc).isoformat()

    conn = _db()
    cur = conn.cursor()

    cur.execute(
        """
        SELECT
          COUNT(*),
          COALESCE(SUM(est_tokens_in),0),
          COALESCE(SUM(est_tokens_out),0)
        FROM usage_events
        WHERE ts >= ?
        """,
        (since_iso,),
    )
    row = cur.fetchone() or (0, 0, 0)
    total_events, tin, tout = row[0], row[1], row[2]

    cur.execute(
        """
        SELECT user_email,
               COUNT(*) as cnt,
               COALESCE(SUM(est_tokens_in),0) as tin,
               COALESCE(SUM(est_tokens_out),0) as tout
        FROM usage_events
        WHERE ts >= ?
        GROUP BY user_email
        ORDER BY (tin+tout) DESC
        LIMIT 50
        """,
        (since_iso,),
    )
    per_user = cur.fetchall() or []

    # Billable documents: count distinct doc_id for "save" actions (preferred),
    # fall back to any conversion activity if no saved events exist.
    cur.execute(
        """
        SELECT COUNT(DISTINCT doc_id)
        FROM usage_events
        WHERE ts >= ?
          AND COALESCE(doc_id,'') <> ''
          AND action IN ('save_version_after_convert','save_version_after_review')
        """,
        (since_iso,),
    )
    billable_docs = int((cur.fetchone() or [0])[0] or 0)

    if billable_docs == 0:
        cur.execute(
            """
            SELECT COUNT(DISTINCT doc_id)
            FROM usage_events
            WHERE ts >= ?
              AND COALESCE(doc_id,'') <> ''
              AND action IN ('llm_convert','llm_reprompt_en')
            """,
            (since_iso,),
        )
        billable_docs = int((cur.fetchone() or [0])[0] or 0)

    conn.close()
    return {

        "since_iso": since_iso,
        "total_events": int(total_events or 0),
        "tokens_in": int(tin or 0),
        "tokens_out": int(tout or 0),
        "tokens_total": int((tin or 0) + (tout or 0)),
        "estimated_cost": money_estimate(int((tin or 0) + (tout or 0))),
        "billable_docs": int(billable_docs or 0),
        "estimated_cost_docs": round(float(BILLING_PRICE_PER_DOC) * float(billable_docs or 0), 2),
        "per_user": per_user,
    }

def analytics_export_csv(days: int = 30) -> bytes:
    since = (datetime.now(timezone.utc).timestamp() - days * 86400)
    since_iso = datetime.fromtimestamp(since, tz=timezone.utc).isoformat()

    conn = _db()
    cur = conn.cursor()
    cur.execute(
        """
        SELECT ts, user_email, action, doc_id, section_id, model,
               prompt_chars, output_chars, est_tokens_in, est_tokens_out, meta_json
        FROM usage_events
        WHERE ts >= ?
        ORDER BY ts DESC
        """,
        (since_iso,),
    )
    rows = cur.fetchall() or []
    conn.close()

    out = io.StringIO()
    w = csv.writer(out)
    w.writerow([
        "ts", "user_email", "action", "doc_id", "section_id", "model",
        "prompt_chars", "output_chars", "est_tokens_in", "est_tokens_out", "meta_json"
    ])
    for r in rows:
        w.writerow(list(r))

    return out.getvalue().encode("utf-8")

# ============================================================
# SQLite storage + ✅ Usage Analytics tables
# ============================================================
def _db() -> sqlite3.Connection:
    os.makedirs(os.path.dirname(SQLITE_PATH) or ".", exist_ok=True)
    conn = sqlite3.connect(SQLITE_PATH, check_same_thread=False)
    conn.execute("PRAGMA journal_mode=WAL;")
    conn.execute("PRAGMA synchronous=NORMAL;")
    return conn

def _db_init() -> None:
    conn = _db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            doc_id TEXT PRIMARY KEY,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            latest_version_id INTEGER
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS versions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_id TEXT NOT NULL,
            version_name TEXT NOT NULL,
            saved_at TEXT NOT NULL,
            snapshot_json TEXT NOT NULL,
            FOREIGN KEY(doc_id) REFERENCES documents(doc_id)
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_versions_doc ON versions(doc_id, id DESC)")

    conn.execute("""
        CREATE TABLE IF NOT EXISTS usage_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ts TEXT NOT NULL,
            user_email TEXT,
            action TEXT NOT NULL,
            doc_id TEXT,
            section_id INTEGER,
            model TEXT,
            prompt_chars INTEGER,
            output_chars INTEGER,
            est_tokens_in INTEGER,
            est_tokens_out INTEGER,
            meta_json TEXT
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_usage_ts ON usage_events(ts DESC)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_usage_user ON usage_events(user_email, ts DESC)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_usage_doc ON usage_events(doc_id, ts DESC)")

    # Audit logs (security + traceability)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS audit_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ts TEXT NOT NULL,
            user_email TEXT,
            role TEXT,
            event TEXT NOT NULL,
            workspace TEXT,
            doc_id TEXT,
            meta_json TEXT
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_audit_ts ON audit_logs(ts DESC)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_audit_user ON audit_logs(user_email, ts DESC)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_audit_role ON audit_logs(role, ts DESC)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_audit_ws ON audit_logs(workspace, ts DESC)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_audit_doc ON audit_logs(doc_id, ts DESC)")

    # Role assignments (admin UI, per-workspace scoping)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS role_assignments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ts TEXT NOT NULL,
            workspace TEXT NOT NULL,      -- '*' means global
            email TEXT NOT NULL,
            role TEXT NOT NULL,           -- admin/reviewer/editor/viewer
            updated_by TEXT
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_roles_email_ws ON role_assignments(email, workspace)")

    # Back-compat: older DBs may not have the new 'role' column.
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(audit_logs)").fetchall()]
        if "role" not in cols:
            conn.execute("ALTER TABLE audit_logs ADD COLUMN role TEXT")
            conn.commit()
    except Exception:
        pass
    conn.commit()
    conn.close()

_db_init()

def log_usage(
    action: str,
    user_email: str = "",
    doc_id: str = "",
    section_id: Optional[int] = None,
    model: str = "",
    prompt_text: str = "",
    output_text: str = "",
    meta: Optional[Dict[str, Any]] = None,
) -> None:
    try:
        prompt_chars = len(prompt_text or "")
        output_chars = len(output_text or "")
        tin = est_tokens_from_text(prompt_text or "")
        tout = est_tokens_from_text(output_text or "")
        meta_json = json.dumps(meta or {}, ensure_ascii=False)

        conn = _db()
        conn.execute(
            """
            INSERT INTO usage_events(ts, user_email, action, doc_id, section_id, model, prompt_chars, output_chars, est_tokens_in, est_tokens_out, meta_json)
            VALUES(?,?,?,?,?,?,?,?,?,?,?)
            """,
            (now_iso(), (user_email or "").lower(), action, doc_id, section_id, model, prompt_chars, output_chars, tin, tout, meta_json),
        )
        conn.commit()
        conn.close()
    except Exception:
        return

# ============================================================
# Audit logging (safe default)
# ============================================================
if AUTH_EMAIL and not st.session_state.get("_logged_login_success"):
    try:
        log_audit(
            event="login_success",
            user_email=AUTH_EMAIL,
            workspace=st.session_state.get("workspace", ""),
            doc_id="",
            meta={"role": st.session_state.get("auth_role")},
        )
        st.session_state["_logged_login_success"] = True
    except Exception:
        pass

def _ensure_doc(doc_id: str) -> None:
    conn = _db()
    cur = conn.cursor()
    cur.execute("SELECT doc_id FROM documents WHERE doc_id = ?", (doc_id,))
    row = cur.fetchone()
    if not row:
        ts = now_iso()
        cur.execute(
            "INSERT INTO documents(doc_id, created_at, updated_at, latest_version_id) VALUES(?,?,?,NULL)",
            (doc_id, ts, ts),
        )
        conn.commit()
    conn.close()

def save_version(doc_id: str, snapshot: Dict[str, Any]) -> str:
    _ensure_doc(doc_id)
    tsname = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    version_name = f"{tsname}.json"
    saved_at = snapshot.get("saved_at") or now_iso()

    raw = json.dumps(snapshot, ensure_ascii=False)
    conn = _db()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO versions(doc_id, version_name, saved_at, snapshot_json) VALUES(?,?,?,?)",
        (doc_id, version_name, saved_at, raw),
    )
    vid = cur.lastrowid
    cur.execute(
        "UPDATE documents SET updated_at = ?, latest_version_id = ? WHERE doc_id = ?",
        (now_iso(), vid, doc_id),
    )
    conn.commit()
    conn.close()
    return version_name

def list_versions(doc_id: str) -> List[str]:
    conn = _db()
    cur = conn.cursor()
    cur.execute("SELECT version_name FROM versions WHERE doc_id = ? ORDER BY id DESC", (doc_id,))
    rows = cur.fetchall()
    conn.close()
    return [r[0] for r in rows] if rows else []

def load_latest(doc_id: str) -> Optional[Dict[str, Any]]:
    conn = _db()
    cur = conn.cursor()
    cur.execute("SELECT latest_version_id FROM documents WHERE doc_id = ?", (doc_id,))
    row = cur.fetchone()
    if not row or not row[0]:
        conn.close()
        return None
    latest_id = row[0]
    cur.execute("SELECT snapshot_json FROM versions WHERE id = ?", (latest_id,))
    row2 = cur.fetchone()
    conn.close()
    if not row2:
        return None
    try:
        return json.loads(row2[0])
    except Exception:
        return None

def load_version(doc_id: str, version_name: str) -> Optional[Dict[str, Any]]:
    conn = _db()
    cur = conn.cursor()
    cur.execute(
        "SELECT snapshot_json FROM versions WHERE doc_id = ? AND version_name = ? ORDER BY id DESC LIMIT 1",
        (doc_id, version_name),
    )
    row = cur.fetchone()
    conn.close()
    if not row:
        return None
    try:
        return json.loads(row[0])
    except Exception:
        return None

def set_latest(doc_id: str, version_name: str) -> bool:
    conn = _db()
    cur = conn.cursor()
    cur.execute(
        "SELECT id FROM versions WHERE doc_id = ? AND version_name = ? ORDER BY id DESC LIMIT 1",
        (doc_id, version_name),
    )
    row = cur.fetchone()
    if not row:
        conn.close()
        return False

    vid = row[0]
    cur.execute(
        "UPDATE documents SET updated_at = ?, latest_version_id = ? WHERE doc_id = ?",
        (now_iso(), vid, doc_id),
    )
    conn.commit()
    conn.close()
    return True

# ============================================================
# Tesseract path (Windows dev helper)
# ============================================================
def find_tesseract_exe_windows() -> Optional[str]:
    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        os.path.join(os.environ.get("LOCALAPPDATA", ""), r"Programs\Tesseract-OCR\tesseract.exe"),
        os.path.join(os.environ.get("LOCALAPPDATA", ""), r"Tesseract-OCR\tesseract.exe"),
    ]
    for p in candidates:
        if p and os.path.isfile(p):
            return p
    return None

_exe = find_tesseract_exe_windows()
if _exe:
    pytesseract.pytesseract.tesseract_cmd = _exe

# ============================================================
# OCR + Confidence
# ============================================================
def pdf_to_images(pdf_bytes: bytes, pages: int, dpi: int = 220) -> List[Image.Image]:
    doc = pdfium.PdfDocument(io.BytesIO(pdf_bytes))
    out: List[Image.Image] = []
    end = min(len(doc), pages)
    scale = dpi / 72
    for i in range(end):
        page = doc[i]
        out.append(page.render(scale=scale).to_pil())
    return out

def _tess_lang_for_detected(code: str) -> str:
    if code == "fr":
        return "fra+osd"
    if code == "en":
        return "eng+osd"
    return "eng+fra+osd"

def ocr_image_with_conf(img: Image.Image, tess_lang: str) -> Dict[str, Any]:
    config = f"--oem 3 --psm 3 -l {tess_lang}"
    text = (pytesseract.image_to_string(img, config=config) or "").strip()

    try:
        data = pytesseract.image_to_data(img, config=config, output_type=pytesseract.Output.DICT)
    except Exception:
        safe_html = "<div style='white-space:pre-wrap'>" + (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") + "</div>"
        return {
            "text": text,
            "html": safe_html,
            "stats": {"words": 0, "avg_conf": 0.0, "low": 0, "med": 0, "high": 0, "conf_available": False},
        }

    n = len(data.get("text", []))
    words_out: List[str] = []
    low = med = high = 0
    conf_sum = 0.0
    conf_count = 0

    def esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    for i in range(n):
        w = (data["text"][i] or "").strip()
        if not w:
            continue
        try:
            conf = float(data["conf"][i])
        except Exception:
            conf = -1.0

        if conf >= 0:
            conf_sum += conf
            conf_count += 1

        if conf < 0:
            words_out.append(esc(w))
            continue

        if conf < OCR_LOW_CONF:
            low += 1
            words_out.append(
                f"<span style='background:rgba(255,0,0,0.18); padding:0 2px; border-radius:4px' title='OCR conf: {conf:.0f}'>"
                f"{esc(w)}</span>"
            )
        elif conf < OCR_MED_CONF:
            med += 1
            words_out.append(
                f"<span style='background:rgba(255,165,0,0.18); padding:0 2px; border-radius:4px' title='OCR conf: {conf:.0f}'>"
                f"{esc(w)}</span>"
            )
        else:
            high += 1
            words_out.append(esc(w))

    avg = (conf_sum / conf_count) if conf_count else 0.0
    html = "<div style='white-space:pre-wrap; line-height:1.6'>" + " ".join(words_out) + "</div>"
    return {
        "text": text,
        "html": html,
        "stats": {"words": conf_count, "avg_conf": round(avg, 2), "low": low, "med": med, "high": high, "conf_available": True},
    }

def extract_pdf(pdf_bytes: bytes, pages: int) -> Tuple[str, Dict[str, Any]]:
    meta: Dict[str, Any] = {
        "pages_requested": pages,
        "text_chars": 0,
        "ocr_chars": 0,
        "ocr_used": False,
        "ocr_failed": False,
        "tesseract_cmd": getattr(pytesseract.pytesseract, "tesseract_cmd", "tesseract"),
        "langdetect_ok": LANGDETECT_OK,
        "langdetect_err": LANGDETECT_ERR,
        "python_exe": sys.executable,
        "detected_lang": "unknown",
        "ocr_detected_lang": "unknown",
        "ocr_conf": {},
    }

    text = ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for i, p in enumerate(reader.pages):
            if i >= pages:
                break
            t = (p.extract_text() or "").strip()
            if t:
                text += t + "\n"
    except Exception:
        text = ""

    text = text.strip()
    meta["text_chars"] = len(text)
    if text:
        meta["detected_lang"] = detect_lang(text)
        return text, meta

    try:
        imgs = pdf_to_images(pdf_bytes, pages, dpi=220)

        first_pass = ocr_image_with_conf(imgs[0], "eng+fra+osd")
        sample = (first_pass.get("text") or "")[:2000]
        ocr_lang = detect_lang(sample)
        meta["ocr_detected_lang"] = ocr_lang

        tess_lang = _tess_lang_for_detected(ocr_lang)

        all_text_parts: List[str] = []
        html_pages: List[str] = []
        stats_pages: List[Dict[str, Any]] = []

        for img in imgs:
            page_res = ocr_image_with_conf(img, tess_lang)
            all_text_parts.append(page_res.get("text", ""))
            html_pages.append(page_res.get("html", ""))
            stats_pages.append(page_res.get("stats", {}))

        ocr_text = "\n\n".join([t for t in all_text_parts if t is not None]).strip()

        meta["ocr_chars"] = len(ocr_text)
        meta["ocr_used"] = True
        meta["detected_lang"] = detect_lang(ocr_text)
        meta["ocr_conf"] = {
            "tess_lang_used": tess_lang,
            "pages": stats_pages,
            "html_pages": html_pages,
            "thresholds": {"low": OCR_LOW_CONF, "med": OCR_MED_CONF},
        }
        return ocr_text, meta
    except Exception:
        meta["ocr_failed"] = True
        return "", meta

def extract_docx(docx_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    d = Document(io.BytesIO(docx_bytes))
    txt = "\n".join(p.text for p in d.paragraphs if p.text.strip()).strip()
    meta = {
        "pages_requested": None,
        "text_chars": len(txt),
        "ocr_chars": 0,
        "ocr_used": False,
        "ocr_failed": False,
        "tesseract_cmd": getattr(pytesseract.pytesseract, "tesseract_cmd", "tesseract"),
        "langdetect_ok": LANGDETECT_OK,
        "langdetect_err": LANGDETECT_ERR,
        "python_exe": sys.executable,
        "detected_lang": detect_lang(txt),
        "ocr_detected_lang": "unknown",
        "ocr_conf": {},
    }
    return txt, meta

# ============================================================
# Headings detection
# ============================================================
HEADING_PAT = re.compile(
    r"^\s*(\d+(\.\d+)*|SECTION|PART|PURPOSE|SCOPE|DEFINITIONS|BACKGROUND|POLICY|RESPONSIBILITIES|APPENDIX|ANNEX)\b",
    re.I
)

def split_by_headings(text: str) -> List[Tuple[str, str]]:
    lines = text.splitlines()
    out: List[Tuple[str, str]] = []
    title: Optional[str] = None
    body: List[str] = []

    for l in lines:
        if HEADING_PAT.match(l.strip()):
            if title and "".join(body).strip():
                out.append((title.strip(), "\n".join(body).strip()))
            title, body = l.strip(), []
        else:
            title = title or "Preamble"
            body.append(l)

    if title and "".join(body).strip():
        out.append((title.strip(), "\n".join(body).strip()))

    return out if out else [("Preamble", text.strip())]

# ============================================================
# Readability scoring (no external deps)
# ============================================================
_VOWELS = "aeiouyàâäæçéèêëîïôöœùûüÿ"

def _count_sentences(text: str) -> int:
    # Rough sentence split on ., !, ?, ; and newlines.
    if not text:
        return 1
    s = re.split(r"[.!?;]+|\n+", text)
    n = sum(1 for part in s if part.strip())
    return max(1, n)

def _count_words(text: str) -> int:
    """Count words in a Unicode-safe way.

    Avoids fragile mojibake character ranges that can raise re.PatternError.
    Matches sequences of Unicode letters (no digits/underscore), with optional apostrophe part.
    """

    if not text:
        return 1
    words = re.findall(r"[^\W\d_]+(?:'[^\W\d_]+)?", text, flags=re.UNICODE)
    return max(1, len(words))

def _count_syllables_word(word: str) -> int:
    """Heuristic syllable counter that is Unicode-safe."""
    w = (word or "").lower()
    if not w:
        return 1

    # Keep only letters (drops punctuation/digits safely, keeps accented letters)
    w = "".join(ch for ch in w if ch.isalpha())
    if not w:
        return 1

    vowels = set("aeiouyàâäæéèêëîïôöœùûüÿ")

    syl = 0
    prev_vowel = False
    for ch in w:
        is_v = ch in vowels
        if is_v and not prev_vowel:
            syl += 1
        prev_vowel = is_v

    # Basic silent-e handling for English
    if w.endswith("e") and not w.endswith(("le", "ye")) and syl > 1:
        syl -= 1

    return max(1, syl)

def _count_syllables(text: str) -> int:
    if not text:
        return 1
    words = re.findall(r"[^\W\d_]+(?:'[^\W\d_]+)?", text, flags=re.UNICODE)
    return sum(_count_syllables_word(w) for w in words) or 1

def flesch_kincaid(text: str) -> float:
    """Flesch–Kincaid Grade Level (English). Lower is easier."""
    words = _count_words(text)
    sentences = _count_sentences(text)
    syllables = _count_syllables(text)
    grade = 0.39 * (words / sentences) + 11.8 * (syllables / words) - 15.59
    # Clamp to sane range for display
    return round(max(0.0, min(20.0, float(grade))), 2)

def french_readability(text: str) -> float:
    """
    French Flesch Reading Ease-style score (higher is easier).
    We keep the same structure: 207 - 1.015*(W/S) - 73.6*(SYL/W).
    """
    words = _count_words(text)
    sentences = _count_sentences(text)
    syllables = _count_syllables(text)
    score = 207.0 - 1.015 * (words / sentences) - 73.6 * (syllables / words)
    return round(float(score), 2)

# ============================================================
# Per-sentence readability helpers (English)
# ============================================================
_SENT_SPLIT_RE = re.compile(r"(?<=[.!?])\s+|\n+")

def split_sentences(text: str) -> List[str]:
    """Split text into sentences (best-effort)."""
    if not text or not text.strip():
        return []
    parts = [p.strip() for p in _SENT_SPLIT_RE.split(text.strip()) if p.strip()]
    return parts

def per_sentence_grades_en(text: str) -> List[Dict[str, Any]]:
    """Return list of {sentence, grade} for English text."""
    out: List[Dict[str, Any]] = []
    for s in split_sentences(text):
        g = flesch_kincaid(s)
        out.append({"sentence": s, "grade": g})
    return out

def worst_sentences_en(text: str, top_n: int = 5) -> List[Dict[str, Any]]:
    rows = per_sentence_grades_en(text)
    rows.sort(key=lambda r: float(r.get("grade", 0.0) or 0.0), reverse=True)
    return rows[: max(0, int(top_n))]

def sentence_heatmap_html_en(text: str, target: float = 8.0) -> str:
    """HTML with sentence-level grade highlighting (English)."""
    rows = per_sentence_grades_en(text)
    if not rows:
        return "<div style='color:#666'>No sentences found.</div>"

    def esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    spans: List[str] = []
    for r in rows:
        s = r["sentence"]
        g = float(r["grade"] or 0.0)
        # simple bands
        if g <= target:
            bg = "rgba(0, 200, 0, 0.12)"
        elif g <= target + 2:
            bg = "rgba(255, 165, 0, 0.16)"
        else:
            bg = "rgba(255, 0, 0, 0.14)"
        spans.append(
            f"<span style='background:{bg}; padding:2px 4px; border-radius:6px; margin:2px; display:inline-block' "
            f"title='Grade: {g:.2f}'>{esc(s)}</span>"
        )
    return "<div style='line-height:1.8'>" + " ".join(spans) + "</div>"

# ============================================================
# LLM (rate-limit safe + caps BEFORE calling) + ✅ usage logging
# ============================================================
def safe_generate(prompt: str, retries: int = LLM_RETRIES):
    delay = 2
    for _ in range(retries):
        try:
            return client.models.generate_content(
                model=LLM_MODEL,
                contents=prompt,
                config={"temperature": 0.1, "max_output_tokens": 550},
            )
        except ClientError:
            time.sleep(delay)
            delay = min(delay * 2, 30)
    return None

def _reprompt_english_to_target(en_text: str, target_grade: float, doc_id: str, section_id: int, user_email: str, max_rounds: int = 6) -> Tuple[str, float, int]:
    """
    Re-prompt ONLY the English text until FK grade <= target_grade.
    Returns (best_text, best_grade, rounds_used).
    """
    best_text = (en_text or "").strip()
    best_grade = flesch_kincaid(best_text) if best_text else 99.0
    rounds = 0
    if not best_text:
        return best_text, best_grade, rounds

    for _ in range(max_rounds):
        if best_grade <= target_grade:
            break

        rounds += 1
        prompt = f"""
Rewrite the ENGLISH text below into plain language at Grade {int(target_grade)} or lower (Flesch–Kincaid).
Rules:
- Keep the same meaning.
- Use short sentences.
- Use common words.
- Active voice.
- Do NOT add headings.
Return JSON ONLY: {{ "en": "..." }}

TEXT:
{best_text}
"""
        resp = safe_generate(prompt)
        raw_out = getattr(resp, "text", "") if resp else ""
        log_usage(
            action="llm_reprompt_en",
            user_email=user_email,
            doc_id=doc_id,
            section_id=section_id,
            model=LLM_MODEL,
            prompt_text=prompt,
            output_text=raw_out,
            meta={"round": rounds, "target_grade": target_grade},
        )

        if not resp or not raw_out.strip():
            continue

        new_en = ""
        try:
            mm = re.search(r"\{.*\}", raw_out, re.S)
            if mm:
                data = json.loads(mm.group())
                new_en = (data.get("en") or "").strip()
        except Exception:
            new_en = ""

        if not new_en:
            new_en = raw_out.strip()

        g = flesch_kincaid(new_en)
        if g < best_grade:
            best_text, best_grade = new_en, g
        else:
            if len(new_en) < len(best_text) and g <= best_grade + 0.25:
                best_text, best_grade = new_en, g

    return best_text, best_grade, rounds

def convert_chunk(text: str, source_lang: str, doc_id: str, section_id: int, user_email: str) -> Dict[str, Any]:
    capped_text, truncated = truncate_words(text, MAX_CHUNK_WORDS)
    best = {"en": "", "fr": "", "grade_en": 99.0, "grade_fr": 0.0, "truncated": truncated, "source_lang": source_lang, "reprompts_en": 0}

    lang_hint = source_lang if source_lang and source_lang != "unknown" else "unknown"
    extra_rules = ""
    if lang_hint == "fr":
        extra_rules = "- Source text is French. Produce English plain-language translation + French plain-language rewrite.\n"
    elif lang_hint == "en":
        extra_rules = "- Source text is English. Produce English plain-language rewrite + French translation.\n"
    else:
        extra_rules = "- Source text language may be mixed/unknown. Preserve meaning; translate as needed.\n"

    # 1) First pass: generate EN+FR together (few tries)
    for attempt in range(3):
        prompt = f"""
Return JSON ONLY: {{ "en": "...", "fr": "..." }}

Rules:
- English Grade ≤ 8 (hard target)
- Short sentences. Active voice.
- Explain terms.
- EN and FR must match meaning.
- No extra text outside JSON.
{extra_rules}
SOURCE_LANG_DETECTED: {lang_hint}

TEXT:
{capped_text}
"""
        r = safe_generate(prompt)
        if not r:
            log_usage(
                action="llm_convert_failed",
                user_email=user_email,
                doc_id=doc_id,
                section_id=section_id,
                model=LLM_MODEL,
                prompt_text=prompt,
                output_text="",
                meta={"attempt": attempt + 1, "reason": "no_response"},
            )
            continue

        raw_out = getattr(r, "text", "") or ""
        log_usage(
            action="llm_convert",
            user_email=user_email,
            doc_id=doc_id,
            section_id=section_id,
            model=LLM_MODEL,
            prompt_text=prompt,
            output_text=raw_out,
            meta={"attempt": attempt + 1, "source_lang": lang_hint, "truncated": truncated},
        )

        try:
            mm = re.search(r"\{.*\}", raw_out, re.S)
            if not mm:
                continue
            data = json.loads(mm.group())
            en = (data.get("en") or "").strip()
            fr = (data.get("fr") or "").strip()
            if not en or not fr:
                continue

            ge = flesch_kincaid(en)
            gf = french_readability(fr)

            if ge < best["grade_en"]:
                best.update({"en": en, "fr": fr, "grade_en": ge, "grade_fr": gf})

            if ge <= 8.0:
                return best
        except Exception:
            continue

    # 2) Hard-enforce: re-prompt ONLY English until compliant (French stays from best pass)
    if best.get("en", "").strip():
        improved_en, improved_grade, rounds = _reprompt_english_to_target(
            best["en"],
            target_grade=8.0,
            doc_id=doc_id,
            section_id=section_id,
            user_email=user_email,
            max_rounds=6,
        )
        best["en"] = improved_en
        best["grade_en"] = improved_grade
        best["reprompts_en"] = rounds

    return best

# ============================================================
# ✅ Hard-enforce Grade ≤ 8 across sections (auto re-simplify EN)
# ============================================================

def _auto_resimplify_sections_to_grade8(
    sections: List[Dict[str, Any]],
    doc_id_scoped: str,
    user_email: str,
    max_rounds_per_section: int = 6,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    """Auto re-simplify EN text for any section above Grade 8.

    Returns (sections, report_rows) where report_rows include title/old/new/rounds.
    Safe: never throws.
    """
    report: List[Dict[str, Any]] = []
    if not sections:
        return sections, report

    for s in sections:
        try:
            en = (s.get("en") or "").strip()
            if not en:
                continue
            old_g = float(s.get("grade_en", flesch_kincaid(en)) or 99.0)
            if old_g <= 8.0:
                continue

            # Use section numeric id if present, else 0
            sid = int(s.get("id", 0) or 0)

            new_en, new_g, rounds = _reprompt_english_to_target(
                en_text=en,
                target_grade=8.0,
                doc_id=doc_id_scoped,
                section_id=sid,
                user_email=user_email,
                max_rounds=max_rounds_per_section,
            )

            # Update section
            s["en"] = new_en
            s["grade_en"] = float(new_g)
            # Keep FR as-is, but refresh FR score (cheap)
            s["grade_fr"] = french_readability(s.get("fr", "") or "")

            report.append({
                "title": s.get("title", f"Section {sid+1}"),
                "old_grade": round(old_g, 2),
                "new_grade": round(float(new_g), 2),
                "rounds": int(rounds),
            })
        except Exception:
            continue

    return sections, report

# ============================================================
# Compliance helpers + Watermarking
# ============================================================
def overall_doc_status(snapshot: Dict[str, Any]) -> str:
    """APPROVED only if every section is approved."""
    secs = snapshot.get("sections", []) or []
    if secs and all((s.get("status") == "approved") for s in secs):
        return "APPROVED"
    return "DRAFT"

def compliance_stats(snapshot: Dict[str, Any]) -> Dict[str, Any]:
    secs = snapshot.get("sections", []) or []
    total = len(secs)
    approved = sum(1 for s in secs if s.get("status") == "approved")
    reviewed = sum(1 for s in secs if s.get("status") == "reviewed")
    draft = sum(1 for s in secs if s.get("status") == "draft")

    grades = []
    over8 = 0
    for s in secs:
        try:
            ge = float(s.get("grade_en", 99.0) or 99.0)
        except Exception:
            ge = 99.0
        grades.append(ge)
        if ge > 8:
            over8 += 1

    truncated = sum(1 for s in secs if s.get("truncated"))

    reviewers = sorted({(s.get("reviewer") or "").strip() for s in secs if (s.get("reviewer") or "").strip()})

    meta = snapshot.get("meta", {}) or {}
    return {
        "total_sections": total,
        "approved": approved,
        "reviewed": reviewed,
        "draft": draft,
        "over_grade8": over8,
        "avg_grade_en": round(sum(grades) / max(1, len(grades)), 2) if grades else 0.0,
        "truncated_sections": truncated,
        "reviewers": reviewers,
        "ocr_used": bool(meta.get("ocr_used", False)),
        "ocr_failed": bool(meta.get("ocr_failed", False)),
        "detected_lang": meta.get("detected_lang", "unknown"),
        "pages_requested": meta.get("pages_requested", ""),
    }

def pdf_watermark(c: canvas.Canvas, text: str, width: float, height: float) -> None:
    """Light diagonal watermark across the page."""
    if not text:
        return
    c.saveState()
    try:
        c.setFillAlpha(0.10)
    except Exception:
        pass
    c.setFont("Helvetica-Bold", 72)
    c.translate(width / 2, height / 2)
    c.rotate(35)
    c.drawCentredString(0, 0, text.upper())
    c.restoreState()

# ============================================================
# Export builders
# ============================================================
def build_docx(snapshot: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("GovCan Plain Language Converter", 1)

    # ✅ Reliable DOCX "watermark": status banner
    status = overall_doc_status(snapshot)
    banner = doc.add_paragraph(f"STATUS: {status}")
    banner.runs[0].bold = True

    meta = snapshot.get("meta", {})
    doc.add_paragraph(f"Document ID: {snapshot.get('doc_id','')}")
    doc.add_paragraph(f"Saved: {snapshot.get('saved_at','')}")
    doc.add_paragraph(f"Detected language: {meta.get('detected_lang', 'unknown')}")
    doc.add_paragraph(f"OCR used: {meta.get('ocr_used', False)} | OCR failed: {meta.get('ocr_failed', False)}")
    doc.add_paragraph(f"Pages requested: {meta.get('pages_requested', '')}")

    ocr_conf = meta.get("ocr_conf", {}) if isinstance(meta.get("ocr_conf", {}), dict) else {}
    if ocr_conf.get("pages"):
        pages = ocr_conf.get("pages", [])
        avg_list = [p.get("avg_conf", 0) for p in pages if isinstance(p, dict)]
        avg = round(sum(avg_list) / max(1, len(avg_list)), 2) if avg_list else 0.0
        doc.add_paragraph(f"OCR avg confidence (approx): {avg}")

    doc.add_paragraph("")

    for s in snapshot.get("sections", []):
        doc.add_heading(s.get("title", "Section"), 2)
        doc.add_paragraph(f"Status: {s.get('status','draft')}")
        doc.add_paragraph(f"Reviewer: {s.get('reviewer','')}")
        doc.add_paragraph(f"Comment: {s.get('comment','')}")
        if s.get("approved_at"):
            doc.add_paragraph(f"Approved at: {s.get('approved_at','')}")

        doc.add_paragraph("")
        doc.add_paragraph("Original").runs[0].bold = True
        doc.add_paragraph(s.get("original", ""))

        doc.add_paragraph("English (Plain Language)").runs[0].bold = True
        doc.add_paragraph(s.get("en", ""))

        doc.add_paragraph("Français (Langage clair)").runs[0].bold = True
        doc.add_paragraph(s.get("fr", ""))

        doc.add_paragraph(f"EN Grade: {s.get('grade_en','')} | FR Readability: {s.get('grade_fr','')}")
        if s.get("truncated"):
            note = doc.add_paragraph(f"NOTE: Section truncated to {MAX_CHUNK_WORDS} words before LLM call.")
            note.runs[0].italic = True

        doc.add_page_break()

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def build_pdf(snapshot: Dict[str, Any]) -> bytes:
    buff = io.BytesIO()
    c = canvas.Canvas(buff, pagesize=letter)
    width, height = letter

    # ✅ Watermark DRAFT vs APPROVED
    status = overall_doc_status(snapshot)
    pdf_watermark(c, status, width, height)

    def draw_wrapped(text: str, x: float, y: float, max_width: float, leading: float = 12) -> float:
        words = text.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test) <= max_width:
                line = test
            else:
                c.drawString(x, y, line)
                y -= leading
                line = w
                if y < 1 * inch:
                    c.showPage()
                    pdf_watermark(c, status, width, height)
                    y = height - 1 * inch
        if line:
            c.drawString(x, y, line)
            y -= leading
        return y

    y = height - 1 * inch
    c.setFont("Helvetica-Bold", 14)
    c.drawString(1 * inch, y, "GovCan Plain Language Converter")
    y -= 18

    c.setFont("Helvetica", 10)
    c.drawString(1 * inch, y, f"Document ID: {snapshot.get('doc_id','')}")
    y -= 14
    c.drawString(1 * inch, y, f"Saved: {snapshot.get('saved_at','')}")
    y -= 14

    meta = snapshot.get("meta", {})
    c.drawString(1 * inch, y, f"Detected language: {meta.get('detected_lang','unknown')}")
    y -= 14
    c.drawString(1 * inch, y, f"OCR used: {meta.get('ocr_used', False)} | OCR failed: {meta.get('ocr_failed', False)} | Pages: {meta.get('pages_requested','')}")
    y -= 18

    for s in snapshot.get("sections", []):
        c.setFont("Helvetica-Bold", 12)
        c.drawString(1 * inch, y, s.get("title", "Section"))
        y -= 16

        c.setFont("Helvetica", 10)
        c.drawString(1 * inch, y, f"Status: {s.get('status','draft')} | Reviewer: {s.get('reviewer','')}")
        y -= 14

        if s.get("comment"):
            y = draw_wrapped(f"Comment: {s.get('comment','')}", 1*inch, y, width - 2*inch, leading=12)

        if s.get("approved_at"):
            y = draw_wrapped(f"Approved at: {s.get('approved_at','')}", 1*inch, y, width - 2*inch, leading=12)

        y -= 8
        c.setFont("Helvetica-Bold", 10)
        c.drawString(1 * inch, y, "Original")
        y -= 12
        c.setFont("Helvetica", 9)
        y = draw_wrapped(s.get("original", ""), 1*inch, y, width - 2*inch, leading=11)

        y -= 6
        c.setFont("Helvetica-Bold", 10)
        c.drawString(1 * inch, y, "English (Plain Language)")
        y -= 12
        c.setFont("Helvetica", 9)
        y = draw_wrapped(s.get("en", ""), 1*inch, y, width - 2*inch, leading=11)

        y -= 6
        c.setFont("Helvetica-Bold", 10)
        c.drawString(1 * inch, y, "Français (Langage clair)")
        y -= 12
        c.setFont("Helvetica", 9)
        y = draw_wrapped(s.get("fr", ""), 1*inch, y, width - 2*inch, leading=11)

        y -= 10
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(1 * inch, y, f"EN Grade: {s.get('grade_en','')} | FR Readability: {s.get('grade_fr','')}")
        y -= 20

        c.showPage()
        pdf_watermark(c, status, width, height)
        y = height - 1 * inch

    c.save()
    return buff.getvalue()

def build_compliance_report_pdf(snapshot: Dict[str, Any]) -> bytes:
    buff = io.BytesIO()
    c = canvas.Canvas(buff, pagesize=letter)
    width, height = letter

    status = overall_doc_status(snapshot)
    pdf_watermark(c, status, width, height)

    stats = compliance_stats(snapshot)

    y = height - 1 * inch
    c.setFont("Helvetica-Bold", 16)
    c.drawString(1 * inch, y, "Compliance Report — GovCan Plain Language Converter")
    y -= 22

    c.setFont("Helvetica", 11)
    c.drawString(1 * inch, y, f"Document ID: {snapshot.get('doc_id','')}")
    y -= 14
    c.drawString(1 * inch, y, f"Version saved at (UTC): {snapshot.get('saved_at','')}")
    y -= 14
    c.drawString(1 * inch, y, f"Overall status: {status}")
    y -= 18

    c.setFont("Helvetica-Bold", 12)
    c.drawString(1 * inch, y, "Summary")
    y -= 16

    c.setFont("Helvetica", 11)

    lines = [
        "Target: English readability at Grade 8 or below (Flesch–Kincaid).",
        "Interpretation: Sections marked OVER exceed target and should be revised.",
        "",
        f"Sections total: {stats['total_sections']}",
        f"Approved: {stats['approved']} | Reviewed: {stats['reviewed']} | Draft: {stats['draft']}",
        f"English grade (avg): {stats['avg_grade_en']} | Sections above Grade 8: {stats['over_grade8']}",
        f"Sections truncated before LLM call: {stats['truncated_sections']}",
        f"Detected language: {stats['detected_lang']} | Pages requested: {stats['pages_requested']}",
        f"OCR used: {stats['ocr_used']} | OCR failed: {stats['ocr_failed']}",
        f"Reviewers: {', '.join(stats['reviewers']) if stats['reviewers'] else '—'}",
    ]
    if stats["ocr_used"]:
        lines.append("Note: OCR text quality may affect readability scores and translation accuracy.")

    for line in lines:
        c.drawString(1 * inch, y, line)
        y -= 14

    y -= 8
    c.setFont("Helvetica-Bold", 12)
    c.drawString(1 * inch, y, "Per-section compliance")
    y -= 16

    c.setFont("Helvetica", 10)
    for s in snapshot.get("sections", []) or []:
        title = (s.get("title") or "Section").strip()
        stt = s.get("status", "draft")
        try:
            ge = float(s.get("grade_en", 99.0) or 99.0)
        except Exception:
            ge = 99.0
        ok = "OK" if ge <= 8 else "OVER"
        trunc = "TRUNC" if s.get("truncated") else ""
        row = f"- {title} | status={stt} | grade_en={ge} ({ok}) {trunc}".strip()

        if y < 1 * inch:
            c.showPage()
            pdf_watermark(c, status, width, height)
            y = height - 1 * inch
            c.setFont("Helvetica", 10)

        c.drawString(1 * inch, y, row[:120])
        y -= 12

    c.save()
    return buff.getvalue()

# ============================================================
# Session state init
# ============================================================
if "doc_id" not in st.session_state:
    st.session_state.doc_id = ""
if "snapshot" not in st.session_state:
    st.session_state.snapshot = None
if "input_text" not in st.session_state:
    st.session_state.input_text = ""
if "extract_meta" not in st.session_state:
    st.session_state.extract_meta = {}
if "last_extract_key" not in st.session_state:
    st.session_state.last_extract_key = ""
if "page_preview_imgs" not in st.session_state:
    st.session_state.page_preview_imgs = []

# ============================================================
# Sidebar
# ============================================================
with st.sidebar:
    st.header("Controls")
    st.caption("Role (locked from Secrets)")
    st.write(f"**{st.session_state.get('auth_role','viewer')}**")
    if AUTH_EMAIL:
        st.caption(f"Signed in as: {AUTH_EMAIL}")
    st.caption(f"Workspace (active): {st.session_state.get('workspace', 'default')}")
    st.caption(f"Workspace (locked): {st.session_state.get('workspace_locked', 'default')}")

    # ============================================================
    # Admin: Role editor (per-workspace) — stored in SQLite
    # ============================================================
    if can("role_admin"):
        with st.expander("Admin: Roles", expanded=False):
            st.caption("Assign roles per workspace. Use '*' for global scope.")
            ws_opt = st.text_input("Workspace scope", value=st.session_state.get("workspace", "default"))
            em_opt = st.text_input("User email", value=AUTH_EMAIL or "")
            rl_opt = st.selectbox("Role", ["admin", "editor", "reviewer", "viewer"], index=0)
            cols = st.columns(2)
            with cols[0]:
                if st.button("Save role", use_container_width=True):
                    _set_role_assignment(ws_opt, em_opt, rl_opt, updated_by=AUTH_EMAIL)
                    log_audit(event="role_assignment_saved", user_email=AUTH_EMAIL, workspace=ws_opt, role=st.session_state.get("auth_role",""), meta={"email": em_opt, "assigned_role": rl_opt})
                    st.success("Saved.")
                    st.rerun()
            with cols[1]:
                if st.button("Grant global admin", use_container_width=True):
                    _set_role_assignment("*", em_opt, "admin", updated_by=AUTH_EMAIL)
                    log_audit(event="role_assignment_saved", user_email=AUTH_EMAIL, workspace="*", role=st.session_state.get("auth_role",""), meta={"email": em_opt, "assigned_role": "admin"})
                    st.success("Granted.")
                    st.rerun()

            st.divider()
            rows = _list_role_assignments(limit=50)
            if rows:
                st.dataframe(rows, use_container_width=True, hide_index=True)
            else:
                st.info("No role assignments yet.")

    if hasattr(st, "logout") and st.button("Logout", use_container_width=True):
        st.logout()

    st.divider()
    st.caption("Storage mode")
    st.success(f"SQLite ✅ ({SQLITE_PATH})")

    st.divider()
    doc_id_in = st.text_input("Document ID", value=st.session_state.doc_id, placeholder="e.g., client-abc-001")

    # ✅ FIX: changing document id should not lock extraction cache
    if doc_id_in != st.session_state.doc_id:
        st.session_state.doc_id = doc_id_in.strip()
        st.session_state.snapshot = None
        st.session_state.last_extract_key = ""  # ✅ IMPORTANT

    c1, c2 = st.columns(2)
    with c1:
        if st.button("Load latest", use_container_width=True):
            if st.session_state.doc_id:
                snap = load_latest(scoped_doc_id(st.session_state.doc_id, st.session_state.workspace))
                st.session_state.snapshot = snap
                st.session_state.last_extract_key = ""
                if snap:
                    st.session_state.input_text = snap.get("source_text", "")
                    st.session_state.extract_meta = snap.get("meta", {})
                st.rerun()
            else:
                st.error("Enter Document ID first.")

    with c2:
        if st.button("New doc", use_container_width=True):
            st.session_state.doc_id = f"doc-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}"
            st.session_state.snapshot = None
            st.session_state.input_text = ""
            st.session_state.extract_meta = {}
            st.session_state.last_extract_key = ""
            st.session_state.page_preview_imgs = []
            st.rerun()

    st.divider()
    if st.session_state.doc_id:
        versions = list_versions(scoped_doc_id(st.session_state.doc_id, st.session_state.workspace))
        if versions:
            vsel = st.selectbox("Version history", versions, index=0)
            if can("rollback") and st.button("Rollback to selected", use_container_width=True):
                ok = set_latest(scoped_doc_id(st.session_state.doc_id, st.session_state.workspace), vsel)
                if ok:
                    snap = load_latest(scoped_doc_id(st.session_state.doc_id, st.session_state.workspace))
                    st.session_state.snapshot = snap
                    st.session_state.last_extract_key = ""
                    if snap:
                        st.session_state.input_text = snap.get("source_text", "")
                        st.session_state.extract_meta = snap.get("meta", {})
                    log_audit(event="rollback", user_email=AUTH_EMAIL, workspace=st.session_state.workspace, doc_id=st.session_state.doc_id, meta={"version": vsel})
                    st.success(f"Rolled back to {vsel}")
                    st.rerun()
                else:
                    st.error("Rollback failed (version not found).")
        else:
            st.caption("No versions yet.")

    # ✅ Admin analytics panel (billing info)
    if can("analytics"):
        st.divider()
        st.subheader("Usage Analytics")
        days = st.slider("Lookback (days)", 1, 180, 30)
        summ = analytics_summary(days=days)
        st.caption(f"Since: {summ['since_iso']}")
        st.write(f"Events: **{summ['total_events']}**")
        st.write(f"Documents processed (billable): **{summ.get('billable_docs', 0)}**")
        st.write(f"Estimated cost (@ ${BILLING_PRICE_PER_DOC:.2f}/doc): **${summ.get('estimated_cost_docs', 0):.2f} CAD**")

        with st.expander("Token diagnostics (admin)"):
            st.write(f"Tokens (in/out): **{summ['tokens_in']} / {summ['tokens_out']}**  | Total: **{summ['tokens_total']}**")
            if BILLING_RATE_PER_1K > 0:
                st.write(f"Token-based estimate (@ {BILLING_RATE_PER_1K}/1K): **{summ['estimated_cost']}**")

        if summ["per_user"]:
            st.caption("Top users (by tokens):")
            for (email, cnt, tin, tout) in summ["per_user"][:10]:
                st.write(f"- {email or '(unknown)'} | events={cnt} | tokens={int(tin)+int(tout)}")

        csv_bytes = analytics_export_csv(days=days)
        st.download_button(
            "Download usage CSV",
            data=csv_bytes,
            file_name=f"usage_events_last_{days}_days.csv",
            mime="text/csv",
            use_container_width=True
        )

        # ðŸ§¾ Admin Audit Logs (security trace)
        st.divider()
        st.subheader("Audit Logs")

        ws_filter_mode = st.selectbox("Workspace scope", ["Current workspace", "All workspaces"], index=0)
        audit_days = st.slider("Audit lookback (days)", 1, 180, 30, key="audit_days")
        q_user = st.text_input("Filter: user email contains", value="", key="audit_user_contains")
        q_role = st.text_input("Filter: role contains", value="", key="audit_role_contains")
        q_event = st.text_input("Filter: event contains", value="", key="audit_event_contains")
        max_rows = st.slider("Rows to show", 50, 500, 200, key="audit_max_rows")

        def audit_query(days: int, workspace: str, all_workspaces: bool, user_contains: str, role_contains: str, event_contains: str, limit: int):
            since = (datetime.now(timezone.utc).timestamp() - days * 86400)
            since_iso = datetime.fromtimestamp(since, tz=timezone.utc).isoformat()
            conn = _db()
            cur = conn.cursor()
            sql = "SELECT ts, user_email, role, event, workspace, doc_id, meta_json FROM audit_logs WHERE ts >= ?"
            params: List[Any] = [since_iso]
            if not all_workspaces:
                sql += " AND workspace = ?"
                params.append(workspace or "default")
            if user_contains.strip():
                sql += " AND lower(user_email) LIKE ?"
                params.append(f"%{user_contains.strip().lower()}%")
            if role_contains.strip():
                sql += " AND lower(role) LIKE ?"
                params.append(f"%{role_contains.strip().lower()}%")
            if event_contains.strip():
                sql += " AND lower(event) LIKE ?"
                params.append(f"%{event_contains.strip().lower()}%")
            sql += " ORDER BY ts DESC LIMIT ?"
            params.append(int(limit))
            cur.execute(sql, tuple(params))
            rows = cur.fetchall() or []
            conn.close()
            return since_iso, rows

        all_ws = (ws_filter_mode == "All workspaces")
        since_iso, rows = audit_query(audit_days, st.session_state.get("workspace","default"), all_ws, q_user, q_role, q_event, max_rows)
        st.caption(f"Showing up to {max_rows} rows since {since_iso}")

        if rows:
            # Render a lightweight table
            st.dataframe(
                [{"ts": r[0], "user": r[1], "role": r[2], "event": r[3], "workspace": r[4], "doc_id": r[5], "meta": r[6]} for r in rows],
                use_container_width=True,
                hide_index=True,
            )

            # CSV export
            out = io.StringIO()
            w = csv.writer(out)
            w.writerow(["ts", "user_email", "role", "event", "workspace", "doc_id", "meta_json"])
            for r in rows:
                w.writerow(list(r))
            st.download_button(
                "Download audit CSV (shown rows)",
                data=out.getvalue().encode("utf-8"),
                file_name=f"audit_logs_{'all' if all_ws else st.session_state.get('workspace','default')}_{audit_days}d.csv",
                mime="text/csv",
                use_container_width=True,
            )
        else:
            st.info("No audit logs found for these filters.")

# ============================================================
# Main UI
# ============================================================
left, right = st.columns(2)

with left:
    up = st.file_uploader("Upload PDF or DOCX", ["pdf", "docx"])
    pages = st.slider("PDF pages to extract/preview", 1, 25, 3)

    preview = st.toggle("Preview pages", value=False)
    show_ocr_conf = st.toggle("Show OCR confidence highlighting (OCR only)", value=False)

    if up:
        b = up.getvalue()
        key = f"{up.name}:{sha12(b)}:pages={pages}"

        should_extract = (key != st.session_state.last_extract_key) or (not st.session_state.input_text.strip())

        if should_extract:
            with st.spinner("Extracting text (OCR if needed)..."):
                if up.name.lower().endswith(".pdf"):
                    txt, meta = extract_pdf(b, pages)
                    if preview:
                        try:
                            st.session_state.page_preview_imgs = pdf_to_images(b, pages, dpi=160)
                        except Exception:
                            st.session_state.page_preview_imgs = []
                else:
                    txt, meta = extract_docx(b)
                    st.session_state.page_preview_imgs = []

                txt, truncated_doc = clamp_text(txt, MAX_DOC_CHARS)
                meta["doc_truncated"] = truncated_doc
                meta["doc_chars_after_cap"] = len(txt)

                st.session_state.last_extract_key = key
                st.session_state.extract_meta = meta
                st.session_state.input_text = txt

    meta = st.session_state.extract_meta or {}
    if meta:
        if meta.get("ocr_failed"):
            st.error("OCR failed ❌ (check tesseract install / language packs)")
        elif meta.get("ocr_used"):
            st.success(f"OCR used ✅ — OCR lang guess: {meta.get('ocr_detected_lang','unknown')} | Detected: {meta.get('detected_lang','unknown')}")
        else:
            st.info(f"Native extraction used ✅ — Detected: {meta.get('detected_lang','unknown')}")

        if meta.get("doc_truncated"):
            st.warning(f"Input capped at {MAX_DOC_CHARS} characters before LLM calls.")

        if not LANGDETECT_OK:
            st.warning("Language detection library not available (langdetect).")
            if LANGDETECT_ERR:
                st.code(LANGDETECT_ERR)
            st.caption(f"Python used by Streamlit: {sys.executable}")

    if show_ocr_conf and meta and meta.get("ocr_used") and isinstance(meta.get("ocr_conf", {}), dict):
        ocr_conf = meta.get("ocr_conf", {})
        html_pages = ocr_conf.get("html_pages", [])
        thr = ocr_conf.get("thresholds", {"low": OCR_LOW_CONF, "med": OCR_MED_CONF})

        st.subheader("OCR confidence highlighting")
        st.caption(f"Red < {thr.get('low')} | Orange < {thr.get('med')} (hover words to see confidence)")
        if html_pages:
            for i, html in enumerate(html_pages, start=1):
                st.markdown(f"**Page {i}**", unsafe_allow_html=False)
                st.markdown(html, unsafe_allow_html=True)
                st.divider()
        else:
            st.info("No OCR confidence data available.")

    if preview and st.session_state.page_preview_imgs:
        st.subheader("Page preview")
        for i, img in enumerate(st.session_state.page_preview_imgs, start=1):
            st.image(img, caption=f"Page {i}", use_container_width=True)

    user_text = st.text_area("Input text", height=320, key="input_text")
    st.caption(f"Chars: {len(user_text)} | Words: {word_count(user_text)}")

    with st.expander("Extraction diagnostics"):
        st.write(meta)

    convert_disabled = (not can("convert")) or (not st.session_state.doc_id) or (not user_text.strip())
    convert_btn = st.button("Convert & Save Version", type="primary", disabled=convert_disabled)

with right:
    if convert_btn:
        safe_text, cut = clamp_text(user_text, MAX_DOC_CHARS)

        doc_lang = detect_lang(safe_text)
        st.session_state.extract_meta = st.session_state.extract_meta or {}
        st.session_state.extract_meta["detected_lang"] = st.session_state.extract_meta.get("detected_lang") or doc_lang

        chunks = split_by_headings(safe_text)

        out_sections: List[Dict[str, Any]] = []
        prog = st.progress(0)
        for i, (title, body) in enumerate(chunks, start=1):
            chunk_lang = detect_lang(body) if len(body) >= 80 else doc_lang
            res = convert_chunk(
                body,
                source_lang=chunk_lang,
                doc_id=scoped_doc_id(st.session_state.doc_id, st.session_state.workspace),
                section_id=i - 1,
                user_email=AUTH_EMAIL,
            )

            out_sections.append({
                "id": i - 1,
                "title": title,
                "original": body,
                "en": res["en"],
                "fr": res["fr"],
                "grade_en": res["grade_en"],
                "grade_fr": res["grade_fr"],
                "truncated": res["truncated"],
                "source_lang": res.get("source_lang", chunk_lang),
                "status": "draft",
                "reviewer": "",
                "comment": "",
                "approved_at": "",
            })
            prog.progress(int(i / max(1, len(chunks)) * 100))
        prog.empty()

        snap = {
            "doc_id": st.session_state.doc_id,
            "saved_at": now_iso(),
            "meta": st.session_state.extract_meta or {},
            "source_text": safe_text,
            "sections": out_sections,
        }
        # ✅ Hard-enforce Grade ≤ 8 (auto re-simplify English before saving)
        doc_id_sc = scoped_doc_id(st.session_state.doc_id, st.session_state.workspace)

        out_sections, _fix_report = _auto_resimplify_sections_to_grade8(
            out_sections,
            doc_id_scoped=doc_id_sc,
            user_email=AUTH_EMAIL,
            max_rounds_per_section=6,
        )

        if _fix_report:
            st.info("Auto-simplified some sections to meet Grade 8 (English).")
            for r in _fix_report[:10]:
                st.write(f"- {r['title']}: {r['old_grade']:.2f} â†’ {r['new_grade']:.2f} (rounds: {r['rounds']})")

        # Final check: block save if still above 8
        over = [
            (s.get("title", "Section"), float(s.get("grade_en", 99.0) or 99.0))
            for s in out_sections
            if float(s.get("grade_en", 99.0) or 99.0) > 8.0
        ]
        if over:
            over_sorted = sorted(over, key=lambda x: x[1], reverse=True)[:10]
            st.error("Grade enforcement failed for one or more sections (must be Grade 8 or below).")
            st.write("Worst sections:")
            for t, g in over_sorted:
                st.write(f"- {t}: {g:.2f}")
            # Keep snapshot in-memory so user can edit, but do NOT save a version
            snap["sections"] = out_sections
            st.session_state.snapshot = snap
            log_audit(
                event="convert_failed_grade",
                user_email=AUTH_EMAIL,
                workspace=st.session_state.workspace,
                doc_id=st.session_state.doc_id,
                meta={"over": over_sorted},
            )
            st.stop()

        snap["sections"] = out_sections
        save_version(doc_id_sc, snap)
        st.session_state.snapshot = snap

        log_usage(
            action="save_version_after_convert",
            user_email=AUTH_EMAIL,
            doc_id=doc_id_sc,
            section_id=None,
            model=LLM_MODEL,
            prompt_text="",
            output_text="",
            meta={"sections": len(out_sections), "doc_chars": len(safe_text), "doc_words": word_count(safe_text)},
        )

        log_audit(
            event="convert_and_save",
            user_email=AUTH_EMAIL,
            workspace=st.session_state.workspace,
            doc_id=st.session_state.doc_id,
            meta={"sections": len(out_sections)},
        )
        st.success("Saved ✅ (SQLite versioned; approvals/comments persist)")

    snap = st.session_state.snapshot
    if not snap:
        st.caption("Upload a document, enter a Document ID, then Convert.")
    else:
        st.subheader("Reviewer workflow + exports")

        # ✅ GovCan compliance banner: Plain-language target met / not met
        _all_grades_ok = True
        _worst = 0.0
        try:
            for _s in (snap.get("sections", []) or []):
                _g = float(_s.get("grade_en", 99.0) or 99.0)
                _worst = max(_worst, _g)
                if _g > 8.0:
                    _all_grades_ok = False
        except Exception:
            _all_grades_ok = False

        if _all_grades_ok:
            st.success("✅ Plain-language target met (English Grade ≤ 8 in all sections).")
        else:
            st.error(f"❌ Plain-language target NOT met (worst section Grade: {_worst:.2f}).")

        st.caption(f"Overall document status: **{overall_doc_status(snap)}**")

        sections = snap.get("sections", [])
        for s in sections:
            st.markdown(f"## {s.get('title','Section')}")

            ge = float(s.get("grade_en", 99.0) or 99.0)
            if ge <= 8:
                st.success(f"EN Grade {ge} ✔")
            else:
                st.warning(f"EN Grade {ge} (above 8)")

            with st.expander("Sentence grade heatmap (English)"):
                st.caption("Green = at/under Grade 8, orange = slightly above, red = high. Hover to see grade.")
                st.markdown(sentence_heatmap_html_en(s.get("en", "")), unsafe_allow_html=True)

                flagged = [r for r in per_sentence_grades_en(s.get("en", "")) if float(r.get("grade", 0.0) or 0.0) > 8.0]
                flagged.sort(key=lambda r: float(r.get("grade", 0.0) or 0.0), reverse=True)
                if flagged:
                    st.subheader("Sentences pushing the grade up")
                    for r in flagged[:8]:
                        st.write(f"- Grade {float(r['grade']):.2f}: {r['sentence']}")
                else:
                    st.success("No sentences above Grade 8 in this section.")

            st.info(f"FR Readability: {s.get('grade_fr', 0.0)}")
            st.caption(f"Source lang detected: {s.get('source_lang','unknown')}")

            if s.get("truncated"):
                st.warning(f"Section truncated to {MAX_CHUNK_WORDS} words before LLM call.")

            if can("edit_outputs"):
                s["en"] = st.text_area("English output", value=s.get("en", ""), key=f"en_{s['id']}", height=140)
                s["fr"] = st.text_area("French output", value=s.get("fr", ""), key=f"fr_{s['id']}", height=140)
                # Recompute readability after edits
                s["grade_en"] = flesch_kincaid(s.get("en", "") or "")
                s["grade_fr"] = french_readability(s.get("fr", "") or "")
            else:
                st.markdown("### English")
                st.write(s.get("en", ""))
                st.markdown("### Français")
                st.write(s.get("fr", ""))

            if can("approve"):
                s["reviewer"] = st.text_input("Reviewer name", value=s.get("reviewer", ""), key=f"rev_{s['id']}")
                s["comment"] = st.text_area("Reviewer comment", value=s.get("comment", ""), key=f"com_{s['id']}", height=90)

                status_options = ["draft", "reviewed", "approved"]
                s["status"] = st.selectbox(
                    "Status",
                    status_options,
                    index=status_options.index(s.get("status", "draft")),
                    key=f"stat_{s['id']}"
                )

                if s["status"] == "approved" and not s.get("approved_at"):
                    s["approved_at"] = now_iso()
                if s["status"] != "approved":
                    s["approved_at"] = ""

            st.divider()

        c1, c2, c3, c4 = st.columns(4)
        with c1:
            if st.button("Save changes as new version", disabled=not st.session_state.doc_id):
                snap["saved_at"] = now_iso()
                snap["sections"] = sections
                # Ensure grades are up-to-date
                for _s in snap["sections"]:
                    _s["grade_en"] = flesch_kincaid(_s.get("en","") or "")
                    _s["grade_fr"] = french_readability(_s.get("fr","") or "")
                # ✅ Hard-enforce Grade ≤ 8 (auto re-simplify EN before saving)
                doc_id_sc = scoped_doc_id(st.session_state.doc_id, st.session_state.workspace)
                snap["sections"], _fix_report2 = _auto_resimplify_sections_to_grade8(
                    snap["sections"],
                    doc_id_scoped=doc_id_sc,
                    user_email=AUTH_EMAIL,
                    max_rounds_per_section=6,
                )

                over2 = [
                    (s.get("title", "Section"), float(s.get("grade_en", 99.0) or 99.0))
                    for s in snap["sections"]
                    if float(s.get("grade_en", 99.0) or 99.0) > 8.0
                ]
                if over2:
                    over_sorted2 = sorted(over2, key=lambda x: x[1], reverse=True)[:10]
                    st.error("Cannot save: plain-language target not met (Grade 8).")
                    st.write("Worst sections:")
                    for t, g in over_sorted2:
                        st.write(f"- {t}: {g:.2f}")
                    log_audit(
                        event="save_failed_grade",
                        user_email=AUTH_EMAIL,
                        workspace=st.session_state.workspace,
                        doc_id=st.session_state.doc_id,
                        meta={"over": over_sorted2},
                    )
                    st.stop()

                if _fix_report2:
                    st.info("Auto-simplified some sections to meet Grade 8 before saving.")

                save_version(scoped_doc_id(st.session_state.doc_id, st.session_state.workspace), snap)
                st.session_state.snapshot = snap

                log_usage(
                    action="save_version_after_review",
                    user_email=AUTH_EMAIL,
                    doc_id=scoped_doc_id(st.session_state.doc_id, st.session_state.workspace),
                    section_id=None,
                    model="",
                    meta={"overall_status": overall_doc_status(snap)},
                )

                log_audit(event="save_review_changes", user_email=AUTH_EMAIL, workspace=st.session_state.workspace, doc_id=st.session_state.doc_id, meta={"overall_status": overall_doc_status(snap)})
                st.success("Saved updated version ✅")

        with c2:
            if can("export"):
                docx_bytes = build_docx(snap)
                st.download_button(
                    "Download DOCX",
                    data=docx_bytes,
                    file_name=f"{safe_filename(st.session_state.doc_id)}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                log_usage(action="export_docx", user_email=AUTH_EMAIL, doc_id=scoped_doc_id(st.session_state.doc_id, st.session_state.workspace), model="", meta={"bytes": len(docx_bytes)})
                log_audit(event="export_docx", user_email=AUTH_EMAIL, workspace=st.session_state.workspace, doc_id=st.session_state.doc_id, meta={"bytes": len(docx_bytes)})

        with c3:
            if can("export"):
                pdf_bytes = build_pdf(snap)
                st.download_button(
                    "Download PDF",
                    data=pdf_bytes,
                    file_name=f"{safe_filename(st.session_state.doc_id)}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
                log_usage(action="export_pdf_full", user_email=AUTH_EMAIL, doc_id=scoped_doc_id(st.session_state.doc_id, st.session_state.workspace), model="", meta={"bytes": len(pdf_bytes)})
                log_audit(event="export_pdf_full", user_email=AUTH_EMAIL, workspace=st.session_state.workspace, doc_id=st.session_state.doc_id, meta={"bytes": len(pdf_bytes)})

        with c4:
            if can("export"):
                comp_pdf = build_compliance_report_pdf(snap)
                st.download_button(
                    "Compliance Report (PDF)",
                    data=comp_pdf,
                    file_name=f"{safe_filename(st.session_state.doc_id)}_compliance.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
                log_usage(action="export_pdf_compliance", user_email=AUTH_EMAIL, doc_id=scoped_doc_id(st.session_state.doc_id, st.session_state.workspace), model="", meta={"bytes": len(comp_pdf)})
                log_audit(event="export_pdf_compliance", user_email=AUTH_EMAIL, workspace=st.session_state.workspace, doc_id=st.session_state.doc_id, meta={"bytes": len(comp_pdf)})

